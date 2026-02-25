"""
人相學書稿工作室 v3
改進：
1. 上傳Word/PDF只處理文字，圖片按章節獨立上傳（解決36MB當機問題）
2. 新增「AI整理逐字稿」功能（貼Whisper稿→Claude API→自動結構化）
3. 章節偵測只抓有【】的主標題
4. 資料庫操作改用先查後寫，避免衝突錯誤
"""

import streamlit as st
import fitz
from docx import Document as DocxDocument
import re
import base64
import datetime
import anthropic
from io import BytesIO
from PIL import Image
from supabase import create_client, Client

# ── 頁面設定 ──────────────────────────────────────────────
st.set_page_config(
    page_title="人相學書稿工作室",
    page_icon="🏮",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Noto+Serif+TC:wght@400;700;900&family=Noto+Sans+TC:wght@300;400;500&display=swap');
html, body, [class*="css"] { font-family: 'Noto Sans TC', sans-serif; }
h1, h2, h3 { font-family: 'Noto Serif TC', serif; }

.title-banner {
    background: linear-gradient(135deg, #8B1A1A 0%, #5C1010 60%, #3A0A0A 100%);
    color: #F5E6C8; padding: 1.8rem 2.5rem; border-radius: 4px;
    margin-bottom: 1.5rem; border-bottom: 4px solid #C9A84C;
}
.title-banner h1 { color: #F5E6C8; font-size: 1.9rem; margin: 0; letter-spacing: 0.12em; }
.title-banner p  { color: #C9A84C; margin: 0.3rem 0 0; font-size: 0.85rem; }

.chapter-card {
    background: white; border-left: 5px solid #8B1A1A;
    border-radius: 0 6px 6px 0; padding: 1rem 1.4rem;
    margin-bottom: 0.8rem; box-shadow: 2px 2px 8px rgba(0,0,0,0.07);
}
.chapter-card h4 { margin: 0 0 0.25rem 0; color: #3A0A0A; font-family: 'Noto Serif TC', serif; }
.chapter-card p  { margin: 0; color: #7A6A5A; font-size: 0.83rem; }

.status-badge { display:inline-block; padding:2px 10px; border-radius:20px; font-size:0.76rem; font-weight:500; margin-left:8px; }
.s-待整理  { background:#F0E8D0; color:#8B6A2A; }
.s-整理中  { background:#D0E8F0; color:#2A5A8B; }
.s-初稿完成 { background:#D0F0DC; color:#2A8B4A; }
.s-修訂中  { background:#F0D8C0; color:#8B4A1A; }
.s-定稿   { background:#3A0A0A; color:#F5E6C8; }

.metric-box { background:white; border:1px solid #E0D5C0; border-top:3px solid #C9A84C; border-radius:4px; padding:1rem; text-align:center; }
.metric-box .num   { font-size:2rem; font-weight:700; color:#8B1A1A; font-family:'Noto Serif TC',serif; }
.metric-box .label { font-size:0.78rem; color:#7A6A5A; margin-top:0.2rem; }
.progress-bg   { background:#E0D5C0; border-radius:4px; height:8px; width:100%; margin-top:4px; }
.progress-fill { height:8px; border-radius:4px; background:linear-gradient(90deg,#8B1A1A,#C9A84C); }
.note-block { background:#FFFBF0; border:1px solid #E8D8A0; border-radius:4px; padding:1rem 1.2rem; font-size:0.87rem; line-height:1.9; white-space:pre-wrap; max-height:480px; overflow-y:auto; color:#2C2010; }
.ai-box { background:#F0F7FF; border:1px solid #B0D0F0; border-radius:4px; padding:1rem 1.2rem; font-size:0.87rem; line-height:1.9; white-space:pre-wrap; max-height:500px; overflow-y:auto; }
</style>
""", unsafe_allow_html=True)

# ── 連線 ─────────────────────────────────────────────────
@st.cache_resource
def get_supabase() -> Client:
    return create_client(st.secrets["SUPABASE_URL"], st.secrets["SUPABASE_KEY"])

@st.cache_resource
def get_claude():
    key = st.secrets.get("ANTHROPIC_API_KEY", "")
    if not key:
        return None
    return anthropic.Anthropic(api_key=key)

supabase = get_supabase()
STATUSES = ['待整理', '整理中', '初稿完成', '修訂中', '定稿']
STATUS_ICON = {'待整理':'⬜','整理中':'🔄','初稿完成':'📝','修訂中':'✏️','定稿':'✅'}

# ── 章節偵測（只抓有【】的主標題，排除「之」子標題）────────
def is_main_chapter(text: str):
    """只抓 一、【形局】 格式，排除 三之一、【xxx】"""
    m = re.match(r'^([一二三四五六七八九十]+)、【([^】]+)】', text.strip())
    if m:
        return True, m.group(1), m.group(2)
    return False, '', ''

def detect_chapters_docx(doc: DocxDocument) -> list[dict]:
    chapters, para_idx = [], 0
    for para in doc.paragraphs:
        matched, num, name = is_main_chapter(para.text.strip())
        if matched:
            if not (chapters and chapters[-1]['name'] == name):
                chapters.append({'num': num, 'name': name,
                                  'start_para': para_idx, 'end_para': None,
                                  'start_page': len(chapters)+1, 'end_page': len(chapters)+1})
        para_idx += 1
    for i, ch in enumerate(chapters):
        ch['end_para'] = chapters[i+1]['start_para']-1 if i+1 < len(chapters) else para_idx
    return chapters

def detect_chapters_pdf(doc) -> list[dict]:
    chapters = []
    for page_idx in range(len(doc)):
        for line in doc[page_idx].get_text("text").strip().split('\n')[:5]:
            matched, num, name = is_main_chapter(line.strip())
            if matched:
                if not (chapters and chapters[-1]['name'] == name):
                    chapters.append({'num': num, 'name': name,
                                      'start_page': page_idx+1, 'end_page': None})
    for i, ch in enumerate(chapters):
        ch['end_page'] = chapters[i+1]['start_page']-1 if i+1 < len(chapters) else len(doc)
    return chapters

# ── 文字提取（只提取文字，不碰圖片）───────────────────────
def extract_text_docx(doc: DocxDocument, start: int, end: int) -> str:
    return '\n'.join(p.text for p in doc.paragraphs[start:end] if p.text.strip())

def extract_text_pdf(doc, start: int, end: int) -> str:
    lines = []
    for pg in range(start-1, min(end, len(doc))):
        t = doc[pg].get_text("text").strip()
        if t:
            lines.append(f"【第{pg+1}頁】\n{t}")
    return '\n\n'.join(lines)

# ── 圖片壓縮工具 ──────────────────────────────────────────
def compress_image(img_bytes: bytes, max_kb=300, max_dim=1200):
    try:
        img = Image.open(BytesIO(img_bytes))
        w, h = img.size
        if w < 50 or h < 50:
            return None, 0, 0
        if w > max_dim or h > max_dim:
            img.thumbnail((max_dim, max_dim), Image.LANCZOS)
            w, h = img.size
        if img.mode in ('RGBA', 'P'):
            img = img.convert('RGB')
        buf, quality = BytesIO(), 75
        img.save(buf, format='JPEG', quality=quality, optimize=True)
        while buf.tell() > max_kb * 1024 and quality > 30:
            quality -= 15
            buf = BytesIO()
            img.save(buf, format='JPEG', quality=quality, optimize=True)
        return base64.b64encode(buf.getvalue()).decode('utf-8'), w, h
    except Exception:
        return None, 0, 0

def extract_images_docx(doc: DocxDocument, start: int, end: int) -> list[dict]:
    images, img_idx = [], 0
    for para_i, para in enumerate(doc.paragraphs[start:end]):
        for run in para.runs:
            for elem in run._element:
                tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
                if tag in ('drawing', 'pict'):
                    r_ids = re.findall(r'r:embed="(rId\d+)"', elem.xml if hasattr(elem, 'xml') else '')
                    for r_id in r_ids:
                        try:
                            rel = doc.part.rels.get(r_id)
                            if rel and 'image' in rel.reltype:
                                b64, w, h = compress_image(rel.target_part.blob)
                                if b64:
                                    images.append({'page': para_i+1, 'index': img_idx,
                                                   'data_b64': b64, 'ext': 'jpeg', 'width': w, 'height': h})
                                    img_idx += 1
                        except Exception:
                            continue
    return images

def extract_images_pdf(doc, start: int, end: int) -> list[dict]:
    images = []
    for pg in range(start-1, min(end, len(doc))):
        for img_idx, img_info in enumerate(doc[pg].get_images(full=True)):
            try:
                base_img = doc.extract_image(img_info[0])
                b64, w, h = compress_image(base_img["image"])
                if b64:
                    images.append({'page': pg+1, 'index': img_idx,
                                   'data_b64': b64, 'ext': 'jpeg', 'width': w, 'height': h})
            except Exception:
                continue
    return images

# ── Supabase 操作（先查後寫，避免 upsert 衝突）─────────────
def db_save_chapter(ch: dict):
    try:
        ex = supabase.table("chapters").select("id").eq("num", ch["num"]).eq("name", ch["name"]).execute()
        if ex.data:
            supabase.table("chapters").update({
                "start_page": ch["start_page"], "end_page": ch["end_page"],
                "last_edit": str(datetime.date.today())
            }).eq("num", ch["num"]).eq("name", ch["name"]).execute()
        else:
            supabase.table("chapters").insert({
                "num": ch["num"], "name": ch["name"],
                "start_page": ch["start_page"], "end_page": ch["end_page"],
                "status": "待整理", "completeness": 0,
                "notes": "", "extra_notes": "", "last_edit": str(datetime.date.today())
            }).execute()
    except Exception as e:
        st.warning(f"章節儲存問題：{e}")

def db_save_text(num: str, name: str, text: str):
    try:
        ex = supabase.table("chapter_content").select("id").eq("chapter_num", num).eq("chapter_name", name).execute()
        if ex.data:
            supabase.table("chapter_content").update({
                "text_content": text, "updated_at": datetime.datetime.utcnow().isoformat()
            }).eq("chapter_num", num).eq("chapter_name", name).execute()
        else:
            supabase.table("chapter_content").insert({
                "chapter_num": num, "chapter_name": name, "text_content": text,
                "updated_at": datetime.datetime.utcnow().isoformat()
            }).execute()
    except Exception as e:
        st.warning(f"文字儲存問題：{e}")

def db_save_images(num: str, name: str, images: list):
    if not images:
        return
    try:
        supabase.table("chapter_images").delete().eq("chapter_num", num).eq("chapter_name", name).execute()
    except Exception:
        pass
    for img in images:
        try:
            supabase.table("chapter_images").insert({
                "chapter_num": num, "chapter_name": name,
                "page_num": img["page"], "img_index": img["index"],
                "data_b64": img["data_b64"], "ext": img["ext"],
                "width": img["width"], "height": img["height"],
            }).execute()
        except Exception:
            continue

def db_update_status(num: str, name: str, status: str, pct: int, notes: str, extra: str):
    try:
        supabase.table("chapters").update({
            "status": status, "completeness": pct,
            "notes": notes, "extra_notes": extra,
            "last_edit": str(datetime.date.today())
        }).eq("num", num).eq("name", name).execute()
    except Exception as e:
        st.error(f"更新失敗：{e}")

def db_load_chapters():
    try:
        return supabase.table("chapters").select("*").order("id").execute().data or []
    except Exception:
        return []

def db_load_text(num: str, name: str) -> str:
    try:
        r = supabase.table("chapter_content").select("text_content").eq("chapter_num", num).eq("chapter_name", name).execute()
        return r.data[0]["text_content"] if r.data else ""
    except Exception:
        return ""

def db_load_images(num: str, name: str) -> list:
    try:
        return supabase.table("chapter_images").select("*").eq("chapter_num", num).eq("chapter_name", name).order("page_num").execute().data or []
    except Exception:
        return []

# ── AI 整理逐字稿 ─────────────────────────────────────────
def ai_organize_transcript(transcript: str, chapter_name: str) -> str:
    """呼叫 Claude API 將 Whisper 逐字稿整理成結構化筆記"""
    client = get_claude()
    prompt = f"""你是一位人相學（面相學）的專業助理。
以下是一段課堂錄音的逐字稿（來自語音辨識，可能有錯字或語氣詞）。
這段內容屬於「{chapter_name}」章節。

請幫我整理成「逐條逐義」的結構化筆記，要求：
1. 每個知識點獨立一條，用數字編號
2. 保留老師的核心論述和斷法，不要刪減重要內容
3. 修正明顯的語音辨識錯字（如相學術語）
4. 若有案例分析，獨立標示為「【案例】」
5. 用繁體中文輸出

逐字稿內容：
{transcript}

請直接輸出整理好的筆記，不要加前言或解釋。"""

    message = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}]
    )
    return message.content[0].text

# ── 工具函數 ──────────────────────────────────────────────
def progress_bar(pct: int) -> str:
    return f'<div class="progress-bg"><div class="progress-fill" style="width:{pct}%"></div></div>'

def render_images(images: list):
    if not images:
        st.caption("此章節無圖片")
        return
    cols = st.columns(min(len(images), 4))
    for i, img in enumerate(images):
        with cols[i % 4]:
            st.image(base64.b64decode(img["data_b64"]),
                     caption=f"第{img['page_num']}頁 ({img['width']}×{img['height']})")

# ══════════════════════════════════════════════════════════
# Sidebar
# ══════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("### 🏮 人相學書稿工作室")

    # ── 上傳筆記（只提取文字）──
    st.markdown("**📂 上傳筆記（文字快速模式）**")
    st.caption("只提取文字和章節，圖片在章節整理頁單獨上傳")
    uploaded = st.file_uploader("Word / PDF", type=["docx", "pdf"], label_visibility="collapsed")

    # 用 session_state 避免重複處理同一個檔案（防止抖動）
    if uploaded:
        file_id = f"{uploaded.name}_{uploaded.size}"
        if st.session_state.get("last_uploaded") != file_id:
            st.session_state["last_uploaded"] = file_id
            file_type = uploaded.name.split(".")[-1].lower()
            with st.spinner("偵測章節並提取文字中..."):
                try:
                    raw = uploaded.read()
                    if file_type == "docx":
                        doc = DocxDocument(BytesIO(raw))
                        chapters = detect_chapters_docx(doc)
                        if not chapters:
                            st.error("找不到章節標題，請確認章節格式為「一、【形局】」")
                        else:
                            bar = st.progress(0)
                            for i, ch in enumerate(chapters):
                                text = extract_text_docx(doc, ch["start_para"], ch["end_para"])
                                db_save_chapter(ch)
                                db_save_text(ch["num"], ch["name"], text)
                                bar.progress((i+1)/len(chapters))
                            st.success(f"✅ 完成！偵測到 {len(chapters)} 個章節")
                    else:
                        doc = fitz.open(stream=raw, filetype="pdf")
                        chapters = detect_chapters_pdf(doc)
                        if not chapters:
                            st.error("找不到章節標題")
                        else:
                            bar = st.progress(0)
                            for i, ch in enumerate(chapters):
                                text = extract_text_pdf(doc, ch["start_page"], ch["end_page"])
                                db_save_chapter(ch)
                                db_save_text(ch["num"], ch["name"], text)
                                bar.progress((i+1)/len(chapters))
                            doc.close()
                            st.success(f"✅ 完成！偵測到 {len(chapters)} 個章節")
                except Exception as e:
                    st.error(f"上傳失敗：{e}")

    st.divider()

    # 整體進度
    all_chapters = db_load_chapters()
    if all_chapters:
        overall = sum(c["completeness"] for c in all_chapters) // len(all_chapters)
        st.markdown(f"""
        <div class="metric-box">
            <div class="num">{overall}%</div>
            <div class="label">整體完成度</div>
            {progress_bar(overall)}
        </div>""", unsafe_allow_html=True)
    else:
        st.info("請先上傳筆記檔案")

    st.divider()
    st.markdown("**📖 功能選單**")
    page = st.radio("", ["🏠 總覽", "📚 章節整理", "🤖 AI整理逐字稿", "📤 匯出書稿"],
                    label_visibility="collapsed")

# ── Banner ────────────────────────────────────────────────
st.markdown("""
<div class="title-banner">
    <h1>🏮 人相學相理精要</h1>
    <p>郭證銂師資班課程筆記 · 書稿整理工作室</p>
</div>""", unsafe_allow_html=True)

all_chapters = db_load_chapters()

# ══════════════════════════════════════════════════════════
# 頁面一：總覽
# ══════════════════════════════════════════════════════════
if page == "🏠 總覽":
    if not all_chapters:
        st.info("👈 請先在左側上傳 Word 或 PDF 筆記檔案")
    else:
        done  = sum(1 for c in all_chapters if c["status"]=="定稿")
        draft = sum(1 for c in all_chapters if c["status"] in ["初稿完成","修訂中"])
        ing   = sum(1 for c in all_chapters if c["status"]=="整理中")

        c1,c2,c3,c4 = st.columns(4)
        for col, num, label in zip([c1,c2,c3,c4],
                                   [len(all_chapters), done, draft, ing],
                                   ["總章節數","已定稿","草稿階段","整理中"]):
            with col:
                st.markdown(f'<div class="metric-box"><div class="num">{num}</div><div class="label">{label}</div></div>', unsafe_allow_html=True)

        st.markdown("---")
        st.markdown("### 📋 各章節狀態")
        for ch in all_chapters:
            pct, status = ch["completeness"], ch["status"]
            col_l, col_r = st.columns([6,1])
            with col_l:
                st.markdown(f"""
                <div class="chapter-card">
                    <h4>{STATUS_ICON.get(status,'❓')} 第{ch['num']}章　{ch['name']}
                        <span class="status-badge s-{status}">{status}</span>
                    </h4>
                    <p>最後編輯：{ch.get('last_edit','—')}</p>
                    {progress_bar(pct)}
                </div>""", unsafe_allow_html=True)
            with col_r:
                st.markdown(f"<div style='text-align:center;padding-top:1.2rem;font-size:1.1rem;color:#8B1A1A;font-weight:700'>{pct}%</div>", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════
# 頁面二：章節整理（含圖片上傳）
# ══════════════════════════════════════════════════════════
elif page == "📚 章節整理":
    st.markdown("### 📚 章節整理")
    if not all_chapters:
        st.info("請先上傳筆記")
    else:
        options = [f"第{c['num']}章：{c['name']}" for c in all_chapters]
        sel = st.selectbox("選擇章節", range(len(all_chapters)), format_func=lambda i: options[i])
        ch = all_chapters[sel]

        st.markdown(f"#### 第{ch['num']}章《{ch['name']}》")
        col_s, col_p = st.columns(2)
        with col_s:
            new_status = st.selectbox("整理狀態", STATUSES, index=STATUSES.index(ch.get("status","待整理")))
        with col_p:
            new_pct = st.slider("完成度", 0, 100, ch.get("completeness",0), 5)
        new_notes = st.text_area("備忘錄（私人，不進書稿）", value=ch.get("notes",""), height=60)

        if st.button("💾 儲存狀態", type="primary"):
            db_update_status(ch["num"], ch["name"], new_status, new_pct, new_notes, ch.get("extra_notes",""))
            st.success("✅ 已更新")
            st.rerun()

        st.divider()

        # 圖片上傳區（按章節獨立上傳，不影響其他章節）
        with st.expander("🖼️ 上傳此章節的圖片（從Word或PDF檔案）", expanded=False):
            st.caption("選擇包含此章節圖片的原始檔案，系統只提取圖片，不影響文字資料")
            img_file = st.file_uploader("選擇檔案", type=["docx","pdf"],
                                        key=f"img_{ch['num']}_{ch['name']}",
                                        label_visibility="collapsed")
            if img_file:
                with st.spinner("提取圖片中..."):
                    try:
                        raw = img_file.read()
                        ftype = img_file.name.split(".")[-1].lower()
                        if ftype == "docx":
                            d = DocxDocument(BytesIO(raw))
                            # 找到對應章節的段落範圍
                            target = next((c for c in detect_chapters_docx(d)
                                           if c["name"] == ch["name"]), None)
                            if target:
                                images = extract_images_docx(d, target["start_para"], target["end_para"])
                            else:
                                # 抓整個檔案的圖片
                                images = extract_images_docx(d, 0, len(d.paragraphs))
                        else:
                            d = fitz.open(stream=raw, filetype="pdf")
                            target = next((c for c in detect_chapters_pdf(d)
                                           if c["name"] == ch["name"]), None)
                            if target:
                                images = extract_images_pdf(d, target["start_page"], target["end_page"])
                            else:
                                images = extract_images_pdf(d, 1, len(d))
                            d.close()

                        db_save_images(ch["num"], ch["name"], images)
                        st.success(f"✅ 已上傳 {len(images)} 張圖片")
                        st.rerun()
                    except Exception as e:
                        st.error(f"圖片提取失敗：{e}")

        st.divider()

        # 顯示文字和圖片
        tab_text, tab_img, tab_extra = st.tabs(["📄 原始文字", "🖼️ 圖片", "📝 補充筆記"])

        with tab_text:
            text = db_load_text(ch["num"], ch["name"])
            if text:
                st.markdown(f'<div class="note-block">{text}</div>', unsafe_allow_html=True)
            else:
                st.caption("尚無文字內容")

        with tab_img:
            images = db_load_images(ch["num"], ch["name"])
            st.caption(f"共 {len(images)} 張圖片")
            render_images(images)

        with tab_extra:
            existing = ch.get("extra_notes", "")
            if existing:
                st.markdown(f'<div class="note-block">{existing}</div>', unsafe_allow_html=True)
                st.divider()
            new_extra = st.text_area("新增補充筆記", height=200,
                                     placeholder="在這裡加入新的課堂補充...")
            col1, col2 = st.columns(2)
            with col1:
                if st.button("➕ 附加", use_container_width=True):
                    if new_extra.strip():
                        today = str(datetime.date.today())
                        merged = existing + f"\n\n---\n【{today}】\n{new_extra.strip()}"
                        db_update_status(ch["num"], ch["name"], ch.get("status","待整理"),
                                         ch.get("completeness",0), ch.get("notes",""), merged)
                        st.success("✅ 已附加")
                        st.rerun()
            with col2:
                if st.button("🔄 覆蓋", use_container_width=True):
                    if new_extra.strip():
                        db_update_status(ch["num"], ch["name"], ch.get("status","待整理"),
                                         ch.get("completeness",0), ch.get("notes",""), new_extra.strip())
                        st.success("✅ 已覆蓋")
                        st.rerun()

# ══════════════════════════════════════════════════════════
# 頁面三：AI 整理逐字稿
# ══════════════════════════════════════════════════════════
elif page == "🤖 AI整理逐字稿":
    st.markdown("### 🤖 AI 整理逐字稿")
    st.caption("將 Whisper 語音辨識的生稿貼入，AI 自動整理成結構化筆記，存入對應章節")

    # 檢查是否有 API 金鑰
    claude_client = get_claude()
    if not claude_client:
        st.warning("🔑 尚未設定 ANTHROPIC_API_KEY，AI整理功能暫不可用。")
        st.markdown("""
        **取得免費 API 金鑰：**
        1. 前往 [console.anthropic.com](https://console.anthropic.com) 註冊
        2. 點 **API Keys** → **Create Key**
        3. 複製金鑰，到 Streamlit Cloud → Settings → Secrets 新增：
           ```
           ANTHROPIC_API_KEY = "sk-ant-..."
           ```
        新帳號有免費額度，足夠整理大量筆記。
        """)
        st.stop()

    if not all_chapters:
        st.info("請先上傳筆記建立章節")
    else:
        options = [f"第{c['num']}章：{c['name']}" for c in all_chapters]
        sel = st.selectbox("選擇要存入的章節", range(len(all_chapters)),
                           format_func=lambda i: options[i])
        ch = all_chapters[sel]

        transcript = st.text_area(
            "貼上 Whisper 逐字稿",
            height=300,
            placeholder="把這次上課的錄音辨識稿貼在這裡...\n\n可以包含口語、語氣詞、錯字，AI會自動清洗整理。"
        )

        col_a, col_b = st.columns([1,2])
        with col_a:
            go = st.button("🤖 開始 AI 整理", type="primary", use_container_width=True,
                           disabled=not transcript.strip())

        if go and transcript.strip():
            with st.spinner(f"AI 正在整理《{ch['name']}》的筆記，約需 30-60 秒..."):
                try:
                    result = ai_organize_transcript(transcript, ch["name"])
                    st.session_state[f"ai_result_{ch['num']}"] = result
                except Exception as e:
                    st.error(f"AI 整理失敗：{e}")

        result_key = f"ai_result_{ch['num']}"
        if result_key in st.session_state:
            result = st.session_state[result_key]
            st.markdown("#### 📋 AI 整理結果（請審閱後再存入）")
            st.markdown(f'<div class="ai-box">{result}</div>', unsafe_allow_html=True)
            st.divider()

            col1, col2, col3 = st.columns(3)
            with col1:
                if st.button("✅ 附加到此章節", type="primary", use_container_width=True):
                    existing = ch.get("extra_notes", "")
                    today = str(datetime.date.today())
                    merged = existing + f"\n\n---\n【{today} AI整理】\n{result}"
                    db_update_status(ch["num"], ch["name"], ch.get("status","整理中"),
                                     ch.get("completeness",0), ch.get("notes",""), merged)
                    del st.session_state[result_key]
                    st.success(f"✅ 已附加到《{ch['name']}》！")
                    st.rerun()
            with col2:
                if st.button("🔄 覆蓋舊補充筆記", use_container_width=True):
                    db_update_status(ch["num"], ch["name"], ch.get("status","整理中"),
                                     ch.get("completeness",0), ch.get("notes",""), result)
                    del st.session_state[result_key]
                    st.success("✅ 已覆蓋")
                    st.rerun()
            with col3:
                if st.button("🗑️ 捨棄結果", use_container_width=True):
                    del st.session_state[result_key]
                    st.rerun()

# ══════════════════════════════════════════════════════════
# 頁面四：匯出書稿
# ══════════════════════════════════════════════════════════
elif page == "📤 匯出書稿":
    st.markdown("### 📤 匯出書稿")
    if not all_chapters:
        st.info("請先上傳筆記")
    else:
        overall = sum(c["completeness"] for c in all_chapters) // len(all_chapters)
        done = sum(1 for c in all_chapters if c["status"]=="定稿")
        st.markdown(f"""
        <div class="metric-box" style="text-align:left;padding:1.2rem 2rem">
            整體完成度：{overall}%　｜　已定稿：{done} / {len(all_chapters)} 章
            {progress_bar(overall)}
        </div>""", unsafe_allow_html=True)

        st.divider()
        col_a, col_b = st.columns(2)

        with col_a:
            st.markdown("#### 📄 完整書稿（Markdown）")
            st.caption("含所有章節文字與補充筆記，可在 Notion / Typora 繼續編輯")
            if st.button("生成書稿", type="primary", use_container_width=True):
                today = datetime.date.today().strftime('%Y年%m月%d日')
                lines = [f"# 人相學相理精要\n\n整理日期：{today}\n\n---\n\n## 目錄\n"]
                for c in all_chapters:
                    lines.append(f"- 第{c['num']}章　{c['name']}　（{c['status']} {c['completeness']}%）")
                lines.append("\n---\n")
                for c in all_chapters:
                    text = db_load_text(c["num"], c["name"])
                    extra = c.get("extra_notes","")
                    lines.append(f"\n## 第{c['num']}章　【{c['name']}】\n")
                    if text:
                        lines.append("### 原始筆記\n" + text)
                    if extra:
                        lines.append("\n### 補充整理\n" + extra)
                    lines.append("\n---")
                manuscript = '\n'.join(lines)
                fname = f"人相學精要_{datetime.date.today().strftime('%Y%m%d')}.md"
                st.download_button("⬇️ 下載書稿 .md", manuscript.encode("utf-8"),
                                   file_name=fname, mime="text/markdown", use_container_width=True)

        with col_b:
            st.markdown("#### 📊 進度報告")
            if st.button("生成進度報告", use_container_width=True):
                lines = [f"# 整理進度報告\n生成：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n"]
                lines.append("| 章節 | 名稱 | 狀態 | 完成度 | 最後編輯 |")
                lines.append("|------|------|------|--------|----------|")
                for c in all_chapters:
                    lines.append(f"| 第{c['num']}章 | {c['name']} | {c['status']} | {c['completeness']}% | {c.get('last_edit','—')} |")
                st.download_button("⬇️ 下載進度報告", '\n'.join(lines).encode("utf-8"),
                                   file_name=f"進度報告_{datetime.date.today()}.md",
                                   mime="text/markdown", use_container_width=True)
